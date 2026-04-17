import re
import unicodedata
import difflib
from docx import Document
import mysql.connector


# =========================================================
# CONFIGURAÇÕES
# =========================================================

ARQUIVO_DOCX = r"C:\Users\William\OneDrive\IPI - Muzambinho\ministerio de TI\relatórios\Livro_atas_2024.docx"

DB_CONFIG = {
    "host": "localhost",
    "user": "root",
    "password": "",
    "database": "szjw_cristaos"
}

ID_IGREJA = 3

# Se True, apaga as presenças das atas deste livro/igreja e recria
LIMPAR_PRESENCAS_ANTES = False

# corte mínimo para aceitar match aproximado
LIMIAR_SIMILARIDADE = 0.86


# =========================================================
# DICIONÁRIO DE ORDINAIS
# =========================================================

ORDINAIS = {
    "primeira": 1,
    "segunda": 2,
    "terceira": 3,
    "quarta": 4,
    "quinta": 5,
    "sexta": 6,
    "sétima": 7,
    "setima": 7,
    "oitava": 8,
    "nona": 9,
    "décima": 10,
    "decima": 10,
    "décima primeira": 11,
    "decima primeira": 11,
    "décima segunda": 12,
    "decima segunda": 12,
    "décima terceira": 13,
    "decima terceira": 13,
}


# =========================================================
# LEITURA DO WORD
# =========================================================

def ler_docx(caminho):
    doc = Document(caminho)
    partes = []

    for p in doc.paragraphs:
        textos = [node.text for node in p._p.xpath('.//w:t') if node.text]
        txt = ''.join(textos).strip()
        if txt:
            partes.append(txt)

    for tabela in doc.tables:
        for row in tabela.rows:
            celulas_linha = []
            for cell in row.cells:
                textos = []
                for p in cell.paragraphs:
                    ts = [node.text for node in p._p.xpath('.//w:t') if node.text]
                    if ts:
                        textos.append(''.join(ts).strip())
                celula_txt = ' '.join([t for t in textos if t])
                if celula_txt:
                    celulas_linha.append(celula_txt)

            if celulas_linha:
                partes.append(" | ".join(celulas_linha))

    texto = "\n".join(partes)
    texto = texto.replace("\xa0", " ")
    texto = re.sub(r"\r", "\n", texto)
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{2,}", "\n\n", texto)
    return texto.strip()


def separar_blocos_atas(texto):
    padrao = re.compile(
        r"(Ata da\s+.*?Reunião do Conselho da Igreja Presbiteriana Independente de Muzambinho do Ano de 2024.*?)(?=Ata da\s+.*?Reunião do Conselho da Igreja Presbiteriana Independente de Muzambinho do Ano de 2024|FOLHA PARA USO EXCLUSIVO|TERMO DE ENCERRAMENTO|\Z)",
        flags=re.IGNORECASE | re.DOTALL
    )
    return [m.group(1).strip() for m in padrao.finditer(texto)]


# =========================================================
# EXTRAÇÃO DOS DADOS DA ATA
# =========================================================

def extrair_numero_reuniao(bloco):
    m = re.search(
        r"Ata da\s+(.+?)\s+Reunião do Conselho",
        bloco,
        flags=re.IGNORECASE | re.DOTALL
    )
    if not m:
        return None

    ordinal = m.group(1).strip().lower()
    ordinal = re.sub(r"\s+", " ", ordinal)

    if ordinal in ORDINAIS:
        return ORDINAIS[ordinal]

    m2 = re.search(r"(\d+)", ordinal)
    if m2:
        return int(m2.group(1))

    return None


def limpar_nome(nome):
    nome = nome.strip()

    # remove títulos e abreviações do começo
    nome = re.sub(
        r'^(presb\.?|presbítero|rev\.?|reverendo|pastor\.?|pr\.?)\s+',
        '',
        nome,
        flags=re.IGNORECASE
    )

    # remove trechos de assinatura
    nome = re.sub(r'_+', '', nome)
    nome = re.sub(r';+', '', nome)

    # espaços extras
    nome = re.sub(r'\s+', ' ', nome).strip()

    return nome


def normalizar(texto):
    texto = texto.lower().strip()
    texto = unicodedata.normalize('NFKD', texto)
    texto = ''.join(c for c in texto if not unicodedata.combining(c))
    texto = re.sub(r'[^a-z0-9\s]', ' ', texto)
    texto = re.sub(r'\s+', ' ', texto).strip()
    return texto


def extrair_nomes_presentes(bloco):
    """
    Extrai os nomes do rodapé/assinaturas.
    """
    linhas = [l.strip() for l in bloco.splitlines() if l.strip()]
    nomes = []

    for linha in linhas:
        # pega linhas que parecem assinatura
        if re.match(r'^(Presb\.?|Rev\.?|Pastor\.?|Pr\.?)\s+', linha, flags=re.IGNORECASE):
            nome = limpar_nome(linha)
            if nome:
                nomes.append(nome)

    # remove duplicados preservando ordem
    nomes_unicos = []
    vistos = set()
    for n in nomes:
        chave = normalizar(n)
        if chave not in vistos:
            vistos.add(chave)
            nomes_unicos.append(n)

    return nomes_unicos


# =========================================================
# BANCO
# =========================================================

def conectar_mysql():
    return mysql.connector.connect(**DB_CONFIG)


def buscar_atas_do_banco(conexao, id_igreja):
    sql = """
        SELECT id_ata, reuniao_numero
        FROM atas
        WHERE id_igreja = %s
    """
    cursor = conexao.cursor(dictionary=True)
    cursor.execute(sql, (id_igreja,))
    rows = cursor.fetchall()
    cursor.close()

    mapa = {}
    for row in rows:
        mapa[row["reuniao_numero"]] = row["id_ata"]
    return mapa


def buscar_membros(conexao):
    sql = """
        SELECT id_membro, nome_do_membro
        FROM membros
        WHERE nome_do_membro IS NOT NULL
          AND nome_do_membro <> ''
    """
    cursor = conexao.cursor(dictionary=True)
    cursor.execute(sql)
    rows = cursor.fetchall()
    cursor.close()
    return rows


def buscar_id_membro(nome_ata, membros):
    nome_ata_norm = normalizar(nome_ata)

    # 1) match exato normalizado
    for membro in membros:
        nome_membro = membro["nome_do_membro"] or ""
        if normalizar(nome_membro) == nome_ata_norm:
            return membro["id_membro"], nome_membro, "exato"

    # 2) contém / contido
    candidatos = []
    for membro in membros:
        nome_membro = membro["nome_do_membro"] or ""
        nome_membro_norm = normalizar(nome_membro)

        if nome_ata_norm in nome_membro_norm or nome_membro_norm in nome_ata_norm:
            score = difflib.SequenceMatcher(None, nome_ata_norm, nome_membro_norm).ratio()
            candidatos.append((score, membro["id_membro"], nome_membro))

    if candidatos:
        candidatos.sort(reverse=True)
        melhor = candidatos[0]
        if melhor[0] >= 0.75:
            return melhor[1], melhor[2], f"contém ({melhor[0]:.2f})"

    # 3) similaridade geral
    melhor_score = 0
    melhor_id = None
    melhor_nome = None

    for membro in membros:
        nome_membro = membro["nome_do_membro"] or ""
        nome_membro_norm = normalizar(nome_membro)
        score = difflib.SequenceMatcher(None, nome_ata_norm, nome_membro_norm).ratio()
        if score > melhor_score:
            melhor_score = score
            melhor_id = membro["id_membro"]
            melhor_nome = nome_membro

    if melhor_score >= LIMIAR_SIMILARIDADE:
        return melhor_id, melhor_nome, f"aproximado ({melhor_score:.2f})"

    return None, None, None


def presenca_ja_existe(conexao, id_ata, id_membro):
    sql = """
        SELECT 1
        FROM presencas_atas
        WHERE id_ata = %s AND id_membro = %s
        LIMIT 1
    """
    cursor = conexao.cursor()
    cursor.execute(sql, (id_ata, id_membro))
    existe = cursor.fetchone() is not None
    cursor.close()
    return existe


def inserir_presenca(conexao, id_ata, id_membro):
    sql = """
        INSERT INTO presencas_atas (id_ata, id_membro)
        VALUES (%s, %s)
    """
    cursor = conexao.cursor()
    cursor.execute(sql, (id_ata, id_membro))
    conexao.commit()
    novo_id = cursor.lastrowid
    cursor.close()
    return novo_id


def limpar_presencas_do_livro(conexao, id_igreja):
    sql = """
        DELETE pa
        FROM presencas_atas pa
        INNER JOIN atas a ON a.id_ata = pa.id_ata
        WHERE a.id_igreja = %s
    """
    cursor = conexao.cursor()
    cursor.execute(sql, (id_igreja,))
    conexao.commit()
    cursor.close()


# =========================================================
# PROCESSAMENTO
# =========================================================

def main():
    print("Lendo arquivo Word...")
    texto = ler_docx(ARQUIVO_DOCX)

    print("Separando atas...")
    atas_word = separar_blocos_atas(texto)
    print(f"Total de atas no Word: {len(atas_word)}")

    if not atas_word:
        print("Nenhuma ata encontrada no arquivo.")
        return

    conexao = conectar_mysql()

    try:
        if LIMPAR_PRESENCAS_ANTES:
            print("Limpando presenças antigas...")
            limpar_presencas_do_livro(conexao, ID_IGREJA)

        mapa_atas = buscar_atas_do_banco(conexao, ID_IGREJA)
        membros = buscar_membros(conexao)

        print(f"Atas no banco para id_igreja={ID_IGREJA}: {len(mapa_atas)}")
        print(f"Membros carregados: {len(membros)}")

        total_inseridas = 0
        total_ignoradas = 0
        total_nao_encontradas = 0

        for i, bloco in enumerate(atas_word, start=1):
            reuniao_numero = extrair_numero_reuniao(bloco)
            nomes = extrair_nomes_presentes(bloco)

            print("-" * 70)
            print(f"Ata {i} / reunião nº {reuniao_numero}")
            print(f"Nomes extraídos: {len(nomes)}")

            if reuniao_numero not in mapa_atas:
                print(">> Ata não encontrada na tabela atas.")
                total_ignoradas += 1
                continue

            id_ata = mapa_atas[reuniao_numero]

            for nome_ata in nomes:
                id_membro, nome_membro_banco, tipo_match = buscar_id_membro(nome_ata, membros)

                if not id_membro:
                    print(f"   [NÃO ENCONTRADO] {nome_ata}")
                    total_nao_encontradas += 1
                    continue

                if presenca_ja_existe(conexao, id_ata, id_membro):
                    print(f"   [JÁ EXISTE] {nome_ata} -> {nome_membro_banco}")
                    continue

                novo_id = inserir_presenca(conexao, id_ata, id_membro)
                print(f"   [OK] {nome_ata} -> {nome_membro_banco} | match={tipo_match} | id_presenca={novo_id}")
                total_inseridas += 1

        print("=" * 70)
        print(f"Presenças inseridas: {total_inseridas}")
        print(f"Atas ignoradas: {total_ignoradas}")
        print(f"Nomes não encontrados: {total_nao_encontradas}")

    except Exception as e:
        print(f"Erro durante a importação: {e}")
        conexao.rollback()

    finally:
        conexao.close()


if __name__ == "__main__":
    main()