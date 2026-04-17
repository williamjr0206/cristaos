import re
from datetime import datetime
from docx import Document
import mysql.connector

ARQUIVO_DOCX = r"C:\Users\William\OneDrive\IPI - Muzambinho\ministerio de TI\relatórios\Livro_atas_2019.docx"

DB_CONFIG = {
    "host": "localhost",
    "user": "root",
    "password": "",
    "database": "szjw_cristaos"
}

ID_IGREJA = 3
NUMERO_LIVRO_FIXO = 13
LIMPAR_TABELA_ANTES = False

ORDINAIS = {
    "primeira": 1,
    "segunda": 2,
    "terceira": 3,
    "quarta": 4,
    "quinta": 5,
    "sexta": 6,
    "sétima": 7,
    "setima": 7,
    "oitava": 8,
    "nona": 9,
    "décima": 10,
    "decima": 10,
    "décima primeira": 11,
    "decima primeira": 11,
    "décima segunda": 12,
    "decima segunda": 12,
    "décima terceira": 13,
    "decima terceira": 13,
}

def ler_docx(caminho):
    doc = Document(caminho)
    partes = []

    for p in doc.paragraphs:
        textos = [node.text for node in p._p.xpath('.//w:t') if node.text]
        txt = ''.join(textos).strip()
        if txt:
            partes.append(txt)

    for tabela in doc.tables:
        for row in tabela.rows:
            celulas_linha = []
            for cell in row.cells:
                textos = []
                for p in cell.paragraphs:
                    ts = [node.text for node in p._p.xpath('.//w:t') if node.text]
                    if ts:
                        textos.append(''.join(ts).strip())
                celula_txt = ' '.join([t for t in textos if t])
                if celula_txt:
                    celulas_linha.append(celula_txt)

            if celulas_linha:
                partes.append(" | ".join(celulas_linha))

    texto = "\n".join(partes)
    texto = texto.replace("\xa0", " ")
    texto = re.sub(r"\r", "\n", texto)
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{2,}", "\n\n", texto)
    return texto.strip()

def separar_blocos_atas(texto):
    padrao = re.compile(
        r"(Ata da\s+.*?Reunião do Conselho da Igreja Presbiteriana Independente de Muzambinho do Ano de 2024.*?)(?=Ata da\s+.*?Reunião do Conselho da Igreja Presbiteriana Independente de Muzambinho do Ano de 2024|FOLHA PARA USO EXCLUSIVO|TERMO DE ENCERRAMENTO|\Z)",
        flags=re.IGNORECASE | re.DOTALL
    )
    return [m.group(1).strip() for m in padrao.finditer(texto)]

def extrair_numero_livro(texto):
    m = re.search(r"Livro de atas número\s+(\d+)", texto, flags=re.IGNORECASE)
    if m:
        return int(m.group(1))
    return NUMERO_LIVRO_FIXO

def extrair_numero_reuniao(bloco):
    m = re.search(
        r"Ata da\s+(.+?)\s+Reunião do Conselho",
        bloco,
        flags=re.IGNORECASE | re.DOTALL
    )
    if not m:
        return None

    ordinal = m.group(1).strip().lower()
    ordinal = re.sub(r"\s+", " ", ordinal)

    if ordinal in ORDINAIS:
        return ORDINAIS[ordinal]

    m2 = re.search(r"(\d+)", ordinal)
    if m2:
        return int(m2.group(1))

    return None

def extrair_data_reuniao(bloco):
    m = re.search(r"No dia\s+(\d{2}/\d{2}/\d{4})", bloco, flags=re.IGNORECASE)
    if not m:
        return None

    dt = datetime.strptime(m.group(1), "%d/%m/%Y")
    return dt.strftime("%Y-%m-%d %H:%M:%S")

def limpar_texto_ata(bloco):
    texto = bloco.strip()
    texto = texto.replace("\xa0", " ")
    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{2,}", "\n\n", texto)
    return texto.strip()

def conectar_mysql():
    return mysql.connector.connect(**DB_CONFIG)

def limpar_tabela_atas(conexao):
    cursor = conexao.cursor()
    cursor.execute("DELETE FROM atas")
    conexao.commit()
    cursor.close()

def inserir_ata(conexao, numero_livro, reuniao_numero, data_reuniao, id_igreja, ata_texto):
    sql = """
        INSERT INTO atas
            (numero_livro, reuniao_numero, data_reuniao, id_igreja, ata_texto)
        VALUES
            (%s, %s, %s, %s, %s)
    """
    valores = (numero_livro, reuniao_numero, data_reuniao, id_igreja, ata_texto)

    cursor = conexao.cursor()
    cursor.execute(sql, valores)
    conexao.commit()
    novo_id = cursor.lastrowid
    cursor.close()
    return novo_id

def main():
    print("Lendo arquivo Word...")
    texto = ler_docx(ARQUIVO_DOCX)

    numero_livro = extrair_numero_livro(texto)
    print(f"Número do livro encontrado: {numero_livro}")

    print("Separando atas...")
    atas = separar_blocos_atas(texto)
    print(f"Total de atas encontradas: {len(atas)}")

    if atas:
        print("\nPrimeiros 250 caracteres da primeira ata encontrada:\n")
        print(atas[0][:250])

    if not atas:
        print("Nenhuma ata encontrada.")
        return

    conexao = conectar_mysql()

    try:
        if LIMPAR_TABELA_ANTES:
            print("Limpando tabela atas...")
            limpar_tabela_atas(conexao)

        total_inseridas = 0

        for i, bloco in enumerate(atas, start=1):
            reuniao_numero = extrair_numero_reuniao(bloco)
            data_reuniao = extrair_data_reuniao(bloco)
            ata_texto = limpar_texto_ata(bloco)

            print("-" * 70)
            print(f"Ata {i}")
            print(f"Reunião número: {reuniao_numero}")
            print(f"Data reunião: {data_reuniao}")
            print(f"Tamanho texto: {len(ata_texto)} caracteres")

            if not reuniao_numero:
                print(">> Ata ignorada: número da reunião não encontrado.")
                continue

            if not data_reuniao:
                print(">> Ata ignorada: data da reunião não encontrada.")
                continue

            novo_id = inserir_ata(
                conexao,
                numero_livro,
                reuniao_numero,
                data_reuniao,
                ID_IGREJA,
                ata_texto
            )

            print(f">> Inserida com id_ata = {novo_id}")
            total_inseridas += 1

        print("=" * 70)
        print(f"Importação concluída. Total inserido: {total_inseridas}")

    except Exception as e:
        print(f"Erro durante a importação: {e}")
        conexao.rollback()

    finally:
        conexao.close()

if __name__ == "__main__":
    main()